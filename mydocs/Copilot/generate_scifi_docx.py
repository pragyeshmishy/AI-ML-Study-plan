#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════════════╗
║  NEXUS-7 DATALOG COMPILER v3.0 — CYBERPUNK TERMINAL EDITION        ║
║  CLASSIFICATION: RESTRICTED // SIGMA-7 CLEARANCE                    ║
║  Converts markdown intelligence to classified .docx HUD format      ║
╚══════════════════════════════════════════════════════════════════════╝

Usage:
    python docs/generate_scifi_docx.py

Reads:  docs/minibook_llm_rag_evaluation_dbt.md
Writes: docs/NEXUS_7_CLASSIFIED_DATALOG.docx
"""

import re
import sys
from pathlib import Path
from lxml import etree

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════════════════════
#  CYBERPUNK HUD PALETTE — Blade Runner / Iron Man aesthetic
# ═══════════════════════════════════════════════════════════════

BG          = '0B0E17'     # deep space navy — page canvas
BG_PANEL    = '0D1117'     # code/data panel background
BG_ELEVATED = '161B22'     # elevated surfaces (table headers)
BG_ACCENT   = '121225'     # blockquote/alert panels

BODY        = 'A8D8D0'     # holographic soft cyan — body text
BRIGHT      = 'F0F6FC'     # high-intensity white — bold emphasis
DIM         = '3D5A6E'     # muted steel — decorative chrome
TEAL        = '80DEEA'     # italic/secondary emphasis

CLR_RED     = 'FF1744'     # alert red — H1, danger accents
CLR_CYAN    = '00E5FF'     # electric cyan — H2, key accents
CLR_AMBER   = 'FFD740'     # amber — H3, warning/highlight
CLR_GREEN   = '7EE787'     # soft neon green — code text
CLR_NEON    = '76FF03'     # hard neon green — inline code
CLR_BLUE    = '1E3A5F'     # deep blue — subtle borders

FONT = 'Consolas'

# ═══════════════════════════════════════════════════════════════
#  LOW-LEVEL XML HELPERS
# ═══════════════════════════════════════════════════════════════

def _mk(tag):
    return OxmlElement(tag)


def _set_background(doc):
    """VML page background — works in Print Layout unlike plain w:background."""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    V = 'urn:schemas-microsoft-com:vml'
    O = 'urn:schemas-microsoft-com:office:office'

    xml = (
        f'<w:background xmlns:w="{W}" xmlns:v="{V}" xmlns:o="{O}" '
        f'w:color="{BG}">'
        f'<v:background id="_x0000_s1025" o:bwmode="white" '
        f'o:targetscreensize="1024,768">'
        f'<v:fill type="frame" color="#{BG}"/>'
        f'</v:background>'
        f'</w:background>'
    )
    doc.element.insert(0, etree.fromstring(xml))
    doc.settings.element.append(_mk('w:displayBackgroundShape'))


def _force_font(run_el, name=FONT):
    rpr = run_el.get_or_add_rPr()
    old = rpr.find(qn('w:rFonts'))
    if old is not None:
        rpr.remove(old)
    rf = _mk('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), name)
    rpr.insert(0, rf)


def _shd(para, fill):
    ppr = para._element.get_or_add_pPr()
    s = _mk('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill)
    ppr.append(s)


def _borders(para, sides_dict):
    """sides_dict: {'bottom': (color, sz), 'left': (color, sz), ...}"""
    ppr = para._element.get_or_add_pPr()
    pbdr = ppr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = _mk('w:pBdr')
        ppr.append(pbdr)
    for side, (color, sz) in sides_dict.items():
        b = _mk(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), color)
        pbdr.append(b)


def _spacing(para, before=0, after=60, line=None):
    ppr = para._element.get_or_add_pPr()
    sp = _mk('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'), str(after))
    if line:
        sp.set(qn('w:line'), str(line))
        sp.set(qn('w:lineRule'), 'auto')
    ppr.append(sp)


def _indent(para, left=0, hanging=0):
    ppr = para._element.get_or_add_pPr()
    ind = _mk('w:ind')
    ind.set(qn('w:left'), str(left))
    if hanging:
        ind.set(qn('w:hanging'), str(hanging))
    ppr.append(ind)


def _cell_shd(cell, fill):
    tcpr = cell._element.get_or_add_tcPr()
    s = _mk('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill)
    tcpr.append(s)


def _cell_margins(cell, v='30', h='60'):
    tcpr = cell._element.get_or_add_tcPr()
    mar = _mk('w:tcMar')
    for side, val in [('top', v), ('bottom', v), ('start', h), ('end', h)]:
        m = _mk(f'w:{side}')
        m.set(qn('w:w'), val)
        m.set(qn('w:type'), 'dxa')
        mar.append(m)
    tcpr.append(mar)


def _tbl_borders(table, color, sz='4'):
    tblpr = table._element.find(qn('w:tblPr'))
    if tblpr is None:
        tblpr = _mk('w:tblPr')
        table._element.insert(0, tblpr)
    old = tblpr.find(qn('w:tblBorders'))
    if old is not None:
        tblpr.remove(old)
    borders = _mk('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = _mk(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tblpr.append(borders)


def _tbl_width_100pct(table):
    tblpr = table._element.find(qn('w:tblPr'))
    if tblpr is None:
        tblpr = _mk('w:tblPr')
        table._element.insert(0, tblpr)
    tw = _mk('w:tblW')
    tw.set(qn('w:w'), '5000')
    tw.set(qn('w:type'), 'pct')
    tblpr.insert(0, tw)


# ═══════════════════════════════════════════════════════════════
#  STYLED RUNS & INLINE PARSING
# ═══════════════════════════════════════════════════════════════

def srun(para, text, color=BODY, size=Pt(10), bold=False, italic=False):
    r = para.add_run(text)
    r.font.name = FONT
    r.font.size = size
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold
    r.italic = italic
    _force_font(r._element)
    return r


def _br(para):
    r = para.add_run()
    r._element.append(_mk('w:br'))


_INLINE = re.compile(
    r'\*\*\*(.+?)\*\*\*'
    r'|\*\*(.+?)\*\*'
    r'|\*(.+?)\*'
    r'|`([^`]+)`'
    r'|\[([^\]]+)\]\([^)]*\)'
)


def _txt(para, text, base_color=BODY, size=Pt(10)):
    last = 0
    for m in _INLINE.finditer(text):
        if m.start() > last:
            srun(para, text[last:m.start()], base_color, size)
        if m.group(1):
            srun(para, m.group(1), BRIGHT, size, True, True)
        elif m.group(2):
            srun(para, m.group(2), BRIGHT, size, True)
        elif m.group(3):
            srun(para, m.group(3), TEAL, size, italic=True)
        elif m.group(4):
            srun(para, m.group(4), CLR_NEON, size)
        elif m.group(5):
            srun(para, m.group(5), CLR_CYAN, size)
        last = m.end()
    if last < len(text):
        srun(para, text[last:], base_color, size)


# ═══════════════════════════════════════════════════════════════
#  BLOCK-LEVEL ELEMENT GENERATORS
# ═══════════════════════════════════════════════════════════════

def mk_header(doc, text, level):
    p = doc.add_paragraph()
    _shd(p, BG)

    if level == 1:
        _spacing(p, before=420, after=180)
        srun(p, '◆ ', CLR_RED, Pt(20), bold=True)
        srun(p, text.upper(), CLR_RED, Pt(26), bold=True)
        _borders(p, {'bottom': (CLR_RED, 12), 'top': (CLR_BLUE, 4)})
    elif level == 2:
        _spacing(p, before=300, after=120)
        srun(p, '▸ ', CLR_CYAN, Pt(14), bold=True)
        srun(p, text.upper(), CLR_CYAN, Pt(17), bold=True)
        _borders(p, {'bottom': (CLR_BLUE, 6)})
    else:
        _spacing(p, before=220, after=90)
        srun(p, '// ', DIM, Pt(11))
        srun(p, text, CLR_AMBER, Pt(12), bold=True)


def mk_code(doc, code_text):
    lines = code_text.split('\n')

    # Top bar
    bar = doc.add_paragraph()
    _shd(bar, BG_ELEVATED)
    _borders(bar, {
        'top':   (CLR_BLUE, 6),
        'left':  (CLR_CYAN, 8),
        'right': (CLR_BLUE, 4),
    })
    _spacing(bar, before=80, after=0)
    _indent(bar, left=170)
    srun(bar, '┌─ DATASTREAM ', DIM, Pt(7))
    srun(bar, '■ ACTIVE', CLR_GREEN, Pt(7), bold=True)

    # Code body
    p = doc.add_paragraph()
    _shd(p, BG_PANEL)
    _borders(p, {
        'left':  (CLR_CYAN, 8),
        'right': (CLR_BLUE, 4),
    })
    _spacing(p, before=0, after=0, line=264)
    _indent(p, left=170)

    for i, ln in enumerate(lines):
        srun(p, ln, CLR_GREEN, Pt(8.5))
        if i < len(lines) - 1:
            _br(p)

    # Bottom bar
    bbar = doc.add_paragraph()
    _shd(bbar, BG_ELEVATED)
    _borders(bbar, {
        'bottom': (CLR_BLUE, 6),
        'left':   (CLR_CYAN, 8),
        'right':  (CLR_BLUE, 4),
    })
    _spacing(bbar, before=0, after=80)
    _indent(bbar, left=170)
    srun(bbar, f'└─ {len(lines)} lines', DIM, Pt(7))


def mk_divider(doc):
    p = doc.add_paragraph()
    _shd(p, BG)
    _spacing(p, before=100, after=100)
    srun(p, '──── ◈ ', DIM, Pt(8))
    srun(p, '━' * 52, CLR_BLUE, Pt(8))
    srun(p, ' ◈ ────', DIM, Pt(8))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def mk_bullet(doc, text):
    p = doc.add_paragraph()
    _shd(p, BG)
    _spacing(p, before=24, after=24)
    _indent(p, left=500, hanging=340)
    srun(p, '  ▸ ', CLR_CYAN, Pt(10), bold=True)
    _txt(p, text, BODY, Pt(10))


def mk_blockquote(doc, lines_list):
    p = doc.add_paragraph()
    _shd(p, BG_ACCENT)
    _borders(p, {
        'left':   (CLR_RED, 18),
        'top':    (CLR_BLUE, 2),
        'bottom': (CLR_BLUE, 2),
    })
    _spacing(p, before=100, after=100)
    _indent(p, left=284)

    for i, line in enumerate(lines_list):
        s = line.strip()
        if s.startswith('- '):
            srun(p, '  ▸ ', CLR_CYAN, Pt(10), bold=True)
            _txt(p, s[2:], CLR_AMBER, Pt(10))
        else:
            _txt(p, s, CLR_AMBER, Pt(10))
        if i < len(lines_list) - 1:
            _br(p)


def mk_paragraph(doc, text):
    p = doc.add_paragraph()
    _shd(p, BG)
    _spacing(p, before=50, after=70)
    _txt(p, text, BODY, Pt(10))


def mk_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < ncols:
            r.append('')

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _tbl_width_100pct(table)

    for ri, rd in enumerate(rows):
        hdr = (ri == 0)
        for ci, ct in enumerate(rd):
            cell = table.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            clean = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', ct.strip())
            if hdr:
                srun(p, clean.upper(), CLR_CYAN, Pt(8.5), bold=True)
                _cell_shd(cell, BG_ELEVATED)
            else:
                _txt(p, ct.strip(), BODY, Pt(8.5))
                _cell_shd(cell, BG_PANEL if ri % 2 == 1 else BG)
            _cell_margins(cell)

    _tbl_borders(table, CLR_BLUE, sz='4')


# ═══════════════════════════════════════════════════════════════
#  HEADER & FOOTER — classification stamps
# ═══════════════════════════════════════════════════════════════

def _setup_header_footer(doc):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = False

    hdr = sec.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _shd(hp, BG)
    _spacing(hp, after=0)
    srun(hp, '▣ ', CLR_RED, Pt(7), bold=True)
    srun(hp, 'RESTRICTED', CLR_RED, Pt(7), bold=True)
    srun(hp, '  //  NEXUS-7  //  SIGMA CLEARANCE  //  ', DIM, Pt(7))
    srun(hp, 'EYES ONLY', CLR_RED, Pt(7), bold=True)
    srun(hp, ' ▣', CLR_RED, Pt(7), bold=True)

    ftr = sec.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _shd(fp, BG)
    _borders(fp, {'top': (CLR_BLUE, 4)})
    _spacing(fp, before=60, after=0)
    srun(fp, '━━  CLASSIFIED DATALOG  ━━  ', DIM, Pt(7))

    # page number field
    run_obj = fp.add_run()
    run_obj.font.name = FONT
    run_obj.font.size = Pt(7)
    run_obj.font.color.rgb = RGBColor.from_string(CLR_CYAN)
    _force_font(run_obj._element)

    fld_begin = _mk('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run_obj._element.append(fld_begin)

    run_instr = fp.add_run()
    run_instr.font.name = FONT
    run_instr.font.size = Pt(7)
    run_instr.font.color.rgb = RGBColor.from_string(CLR_CYAN)
    _force_font(run_instr._element)
    instr = _mk('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run_instr._element.append(instr)

    run_end = fp.add_run()
    run_end.font.name = FONT
    run_end.font.size = Pt(7)
    run_end.font.color.rgb = RGBColor.from_string(CLR_CYAN)
    _force_font(run_end._element)
    fld_end = _mk('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run_end._element.append(fld_end)

    srun(fp, ' / ', DIM, Pt(7))

    # total pages field
    run_obj2 = fp.add_run()
    run_obj2.font.name = FONT
    run_obj2.font.size = Pt(7)
    run_obj2.font.color.rgb = RGBColor.from_string(CLR_CYAN)
    _force_font(run_obj2._element)
    fld_begin2 = _mk('w:fldChar')
    fld_begin2.set(qn('w:fldCharType'), 'begin')
    run_obj2._element.append(fld_begin2)

    run_instr2 = fp.add_run()
    run_instr2.font.name = FONT
    run_instr2.font.size = Pt(7)
    run_instr2.font.color.rgb = RGBColor.from_string(CLR_CYAN)
    _force_font(run_instr2._element)
    instr2 = _mk('w:instrText')
    instr2.set(qn('xml:space'), 'preserve')
    instr2.text = ' NUMPAGES '
    run_instr2._element.append(instr2)

    run_end2 = fp.add_run()
    run_end2.font.name = FONT
    run_end2.font.size = Pt(7)
    run_end2.font.color.rgb = RGBColor.from_string(CLR_CYAN)
    _force_font(run_end2._element)
    fld_end2 = _mk('w:fldChar')
    fld_end2.set(qn('w:fldCharType'), 'end')
    run_end2._element.append(fld_end2)

    srun(fp, '  ━━  NEXUS-7  ━━', DIM, Pt(7))


# ═══════════════════════════════════════════════════════════════
#  MAIN PARSER & GENERATOR
# ═══════════════════════════════════════════════════════════════

def generate(md, out):
    doc = Document()

    # Page background
    _set_background(doc)

    # Page geometry
    for sec in doc.sections:
        sec.top_margin    = Inches(0.55)
        sec.bottom_margin = Inches(0.55)
        sec.left_margin   = Inches(0.6)
        sec.right_margin  = Inches(0.6)

    # Default style
    ns = doc.styles['Normal']
    ns.font.name = FONT
    ns.font.size = Pt(10)
    ns.font.color.rgb = RGBColor.from_string(BODY)
    rpr = ns.element.find(qn('w:rPr'))
    if rpr is None:
        rpr = _mk('w:rPr')
        ns.element.append(rpr)
    old_rf = rpr.find(qn('w:rFonts'))
    if old_rf is not None:
        rpr.remove(old_rf)
    rf = _mk('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), FONT)
    rpr.insert(0, rf)

    # Paragraph defaults — zero spacing to minimize white gaps
    ppr_default = ns.element.find(qn('w:pPr'))
    if ppr_default is None:
        ppr_default = _mk('w:pPr')
        ns.element.append(ppr_default)
    sp_default = _mk('w:spacing')
    sp_default.set(qn('w:before'), '0')
    sp_default.set(qn('w:after'), '0')
    sp_default.set(qn('w:line'), '276')
    sp_default.set(qn('w:lineRule'), 'auto')
    ppr_default.append(sp_default)

    # Remove default empty paragraph
    body = doc.element.find(qn('w:body'))
    dp = body.find(qn('w:p'))
    if dp is not None:
        body.remove(dp)

    # Metadata
    doc.core_properties.title   = 'NEXUS-7 CLASSIFIED DATALOG'
    doc.core_properties.subject = 'LLM Evaluation // RAG // Re-Ranking // AWS DBT // Prompt Security'
    doc.core_properties.author  = 'NEXUS MAINFRAME // DC-COPILOT DIVISION'

    # Header & footer
    _setup_header_footer(doc)

    # Parse markdown
    lines = md.split('\n')
    n = len(lines)
    i = 0
    st = dict(h=0, p=0, c=0, t=0, b=0, q=0, d=0)

    while i < n:
        line = lines[i]
        s = line.strip()

        # code fence
        if s.startswith('```'):
            cl = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                cl.append(lines[i])
                i += 1
            if i < n:
                i += 1
            if cl:
                mk_code(doc, '\n'.join(cl))
                st['c'] += 1
            continue

        # header
        hm = re.match(r'^(#{1,3})\s+(.*)', line)
        if hm:
            mk_header(doc, hm.group(2).strip(), len(hm.group(1)))
            st['h'] += 1
            i += 1
            continue

        # hr
        if re.match(r'^-{3,}\s*$', s) or re.match(r'^\*{3,}\s*$', s):
            mk_divider(doc)
            st['d'] += 1
            i += 1
            continue

        # table
        if s.startswith('|') and '|' in s[1:]:
            raw = []
            while i < n and lines[i].strip().startswith('|'):
                raw.append(lines[i])
                i += 1
            parsed = []
            for rl in raw:
                cells = [c.strip() for c in rl.strip('|').split('|')]
                if cells and all(re.match(r'^[-:]+$', c) for c in cells if c):
                    continue
                if cells:
                    parsed.append(cells)
            if parsed:
                mk_table(doc, parsed)
                st['t'] += 1
            continue

        # blockquote
        if s.startswith('>'):
            ql = []
            while i < n and lines[i].strip().startswith('>'):
                ql.append(re.sub(r'^>\s?', '', lines[i]))
                i += 1
            mk_blockquote(doc, ql)
            st['q'] += 1
            continue

        # bullet
        if re.match(r'^\s*[-*+]\s', line):
            mk_bullet(doc, re.sub(r'^\s*[-*+]\s', '', line))
            st['b'] += 1
            i += 1
            continue

        # empty
        if not s:
            i += 1
            continue

        # paragraph
        parts = [line]
        i += 1
        while i < n:
            nx = lines[i]
            ns2 = nx.strip()
            if (not ns2
                or ns2.startswith('#')
                or ns2.startswith('```')
                or ns2.startswith('>')
                or ns2.startswith('|')
                or re.match(r'^-{3,}\s*$', ns2)
                or re.match(r'^\*{3,}\s*$', ns2)
                or re.match(r'^\s*[-*+]\s', nx)):
                break
            parts.append(nx)
            i += 1
        mk_paragraph(doc, ' '.join(parts))
        st['p'] += 1

    doc.save(str(out))
    return st


# ═══════════════════════════════════════════════════════════════
#  ENTRY POINT
# ═══════════════════════════════════════════════════════════════

if __name__ == '__main__':
    base = Path(__file__).resolve().parent
    inp = base / 'minibook_llm_rag_evaluation_dbt.md'
    out = base / 'NEXUS_7_CLASSIFIED_DATALOG.docx'

    if not inp.exists():
        print(f'[NEXUS-7] ERROR: Source not found: {inp}')
        sys.exit(1)

    md = inp.read_text(encoding='utf-8')
    print(f'[NEXUS-7] ▓▓▓ CYBERPUNK DATALOG COMPILER v3.0 ▓▓▓')
    print(f'[NEXUS-7] Compiling {len(md.splitlines()):,} lines of classified intel...')

    st = generate(md, out)

    labels = dict(h='Headers', p='Paragraphs', c='Code Blocks',
                  t='Tables', b='Bullets', q='Blockquotes', d='Dividers')
    print(f'[NEXUS-7] Output → {out.name}')
    print(f'[NEXUS-7] ┌──────────────────────────────────┐')
    for k, v in st.items():
        print(f'[NEXUS-7] │  {labels[k]:12s}  {v:>6,}             │')
    print(f'[NEXUS-7] └──────────────────────────────────┘')
    print(f'[NEXUS-7] ▓▓▓ TRANSMISSION COMPLETE ▓▓▓')
